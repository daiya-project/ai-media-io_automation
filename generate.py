import copy
import re
from pathlib import Path
from collections import OrderedDict
import openpyxl
from docx import Document
from docx.oxml.ns import qn

COMMON_FIELDS = ["client_name", "client_address", "client_email", "client_manager", "gross_rate"]
WIDGET_FIELDS = ["service", "service_name", "widget_name", "value", "date_start"]


def read_excel_data(excel_path: Path) -> OrderedDict:
    """엑셀 파일을 읽어 client_name 기준으로 그룹핑하여 반환한다."""
    wb = openpyxl.load_workbook(excel_path)
    ws = wb.active

    headers = [cell.value for cell in ws[1]]
    groups = OrderedDict()

    for row in ws.iter_rows(min_row=2, values_only=True):
        row_dict = dict(zip(headers, row))
        name = str(row_dict["client_name"])

        if name not in groups:
            groups[name] = {
                field: str(row_dict[field] or "") for field in COMMON_FIELDS
            }
            groups[name]["widgets"] = []

        widget = {field: str(row_dict[field] or "") for field in WIDGET_FIELDS}
        groups[name]["widgets"].append(widget)

    return groups


def _replace_in_paragraph(paragraph, replacements: dict):
    """paragraph 내 {변수}를 치환한다. run이 분할된 경우를 처리."""
    full_text = paragraph.text
    if not any(key in full_text for key in replacements):
        return

    for key, value in replacements.items():
        full_text = full_text.replace(key, value)

    # 기존 run의 서식을 보존하면서 텍스트 교체
    if paragraph.runs:
        paragraph.runs[0].text = full_text
        for run in paragraph.runs[1:]:
            run.text = ""


def _replace_in_table(table, replacements: dict):
    """테이블 내 모든 셀의 변수를 치환한다."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)


def _build_widget_rows(table, widgets: list):
    """Table 1의 템플릿 행(Row 1)을 기준으로 위젯 수만큼 행을 생성한다."""
    # Row 0은 헤더, Row 1은 템플릿 행
    template_row = table.rows[1]._tr

    # 기존 데이터 행(Row 1~) 제거
    for row in table.rows[1:]:
        table._tbl.remove(row._tr)

    # 위젯 데이터로 새 행 생성
    for widget in widgets:
        new_row = copy.deepcopy(template_row)
        cells = new_row.findall(qn('w:tc'))
        field_map = [
            widget["service"],
            widget["service_name"],
            widget["widget_name"],
            widget["value"],
            widget["date_start"],
            "-",
        ]
        for cell_elem, value in zip(cells, field_map):
            for p in cell_elem.findall(qn('w:p')):
                runs = p.findall(qn('w:r'))
                if not runs:
                    continue
                # 첫 번째 run의 텍스트를 설정
                t_elems = runs[0].findall(qn('w:t'))
                if t_elems:
                    t_elems[0].text = value
                    # 첫 번째 run의 나머지 <w:t> 제거
                    for t in t_elems[1:]:
                        runs[0].remove(t)
                # 나머지 run 모두 제거 (서식 깨짐 방지)
                for r in runs[1:]:
                    p.remove(r)
                break

        table._tbl.append(new_row)


def generate_document(template_path: Path, data: dict, output_path: Path):
    """템플릿을 복제하고 변수를 치환하여 새 docx를 생성한다."""
    doc = Document(template_path)

    replacements = {
        "{client_name}": data["client_name"],
        "{client_address}": data["client_address"],
        "{client_email}": data["client_email"],
        "{client_manager}": data["client_manager"],
        "{gross_rate}": data["gross_rate"],
    }

    # 문단 치환
    for paragraph in doc.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    # 테이블 치환 (Table 0: Contact, Table 2: 조건, Table 3: 서명)
    for i in [0, 2, 3]:
        _replace_in_table(doc.tables[i], replacements)

    # Table 1: 위젯 행 동적 생성
    _build_widget_rows(doc.tables[1], data["widgets"])

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
