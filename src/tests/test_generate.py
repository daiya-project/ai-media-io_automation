import sys
import pytest
from pathlib import Path
from docx import Document

# src/ 디렉토리를 import 경로에 추가
SRC_DIR = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(SRC_DIR))

TEMPLATE = SRC_DIR / "io-sample.docx"
SAMPLE_DATA = SRC_DIR.parent / "data" / "sample.xlsx"


def test_read_and_group_excel():
    """엑셀 파일을 읽어 client_name 기준으로 그룹핑한다."""
    from generate import read_excel_data

    groups = read_excel_data(SAMPLE_DATA)

    assert len(groups) == 4  # 조선일보, 인벤, 뽐뿌, 매일경제

    group = groups["조선일보"]
    assert group["client_name"] == "조선일보"
    assert group["client_address"] == "서울시 중구 세종대로 21길 30"
    assert group["client_email"] == "ad@chosun.com"
    assert group["client_manager"] == "박지현"
    assert group["gross_rate"] == "55%"
    assert len(group["widgets"]) == 3

    widget0 = group["widgets"][0]
    assert widget0["service"] == "조선일보"
    assert widget0["service_name"] == "chosun.com"
    assert widget0["widget_name"] == "메인 하단 위젯"
    assert widget0["value"] == "CPM 1,500원"
    assert widget0["date_start"] == "2026-04-01"

    assert len(groups["뽐뿌"]["widgets"]) == 1
    assert len(groups["매일경제"]["widgets"]) == 4


def test_replace_simple_variables(tmp_path):
    """docx 내 단순 변수({client_name} 등)를 치환한다."""
    from generate import generate_document

    data = {
        "client_name": "테스트매체A",
        "client_address": "서울시 강남구 역삼동 123",
        "client_email": "a@test.com",
        "client_manager": "김철수",
        "gross_rate": "50%",
        "widgets": [
            {"service": "매체A", "service_name": "a.com", "widget_name": "위젯1",
             "value": "CPM 1,000원", "date_start": "2026-04-01"},
        ],
    }

    output_path = tmp_path / "test-output.docx"
    generate_document(TEMPLATE, data, output_path)

    doc = Document(output_path)

    table0 = doc.tables[0]
    assert "테스트매체A" in table0.rows[1].cells[4].text

    table2 = doc.tables[2]
    assert "50%" in table2.rows[0].cells[1].text

    table3 = doc.tables[3]
    assert "김철수" in table3.rows[1].cells[3].text


def test_dynamic_widget_rows(tmp_path):
    """위젯 수에 따라 Table 1의 행이 동적으로 생성된다."""
    from generate import generate_document

    data = {
        "client_name": "테스트매체",
        "client_address": "주소",
        "client_email": "e@e.com",
        "client_manager": "매니저",
        "gross_rate": "40%",
        "widgets": [
            {"service": f"매체{i}", "service_name": f"s{i}.com",
             "widget_name": f"위젯{i}", "value": f"CPM {i}00원",
             "date_start": "2026-04-01"}
            for i in range(5)
        ],
    }

    output_path = tmp_path / "test-dynamic.docx"
    generate_document(TEMPLATE, data, output_path)

    doc = Document(output_path)
    table1 = doc.tables[1]

    assert len(table1.rows) == 6

    for i in range(5):
        assert f"매체{i}" in table1.rows[i + 1].cells[0].text


def test_read_csv_data(tmp_path):
    """CSV 파일을 읽어 client_name 기준으로 그룹핑한다."""
    from generate import read_input_data

    csv_file = tmp_path / "test.csv"
    csv_file.write_text(
        "client_name,client_address,client_email,client_manager,gross_rate,"
        "service,service_name,widget_name,value,date_start\n"
        "A사,서울시 강남구,a@a.com,김철수,50%,매체A,a.com,위젯1,CPM 1000원,2026-04-01\n"
        "A사,서울시 강남구,a@a.com,김철수,50%,매체A,a.com,위젯2,CPM 800원,2026-04-01\n"
        "B사,부산시 해운대,b@b.com,이영희,45%,매체B,b.com,위젯3,CPM 900원,2026-05-01\n",
        encoding="utf-8-sig",
    )

    groups = read_input_data(csv_file)

    assert len(groups) == 2
    assert groups["A사"]["client_address"] == "서울시 강남구"
    assert len(groups["A사"]["widgets"]) == 2
    assert len(groups["B사"]["widgets"]) == 1
    assert groups["B사"]["widgets"][0]["widget_name"] == "위젯3"
