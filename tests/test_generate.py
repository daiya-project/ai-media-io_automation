import pytest
from pathlib import Path
from docx import Document

def test_read_and_group_excel():
    """엑셀 파일을 읽어 client_name 기준으로 그룹핑한다."""
    from generate import read_excel_data

    groups = read_excel_data(Path("sample-data.xlsx"))

    assert len(groups) == 2  # 테스트매체A, 테스트매체B

    group_a = groups["테스트매체A"]
    assert group_a["client_name"] == "테스트매체A"
    assert group_a["client_address"] == "서울시 강남구 역삼동 123"
    assert group_a["client_email"] == "a@test.com"
    assert group_a["client_manager"] == "김철수"
    assert group_a["gross_rate"] == "50%"
    assert len(group_a["widgets"]) == 3

    widget0 = group_a["widgets"][0]
    assert widget0["service"] == "매체A-1"
    assert widget0["service_name"] == "a1.example.com"
    assert widget0["widget_name"] == "위젯Alpha"
    assert widget0["value"] == "CPM 1,000원"
    assert widget0["date_start"] == "2026-04-01"

    group_b = groups["테스트매체B"]
    assert len(group_b["widgets"]) == 1


def test_replace_simple_variables(tmp_path):
    """docx 내 단순 변수({client_name} 등)를 치환한다."""
    from generate import generate_document

    data = {
        "client_name": "테스트매체A",
        "client_address": "서울시 강남구 역삼동 123",
        "client_email": "a@test.com",
        "client_manager": "김철수",
        "gross_rate": "50%",
        "widgets": [
            {"service": "매체A", "service_name": "a.com", "widget_name": "위젯1",
             "value": "CPM 1,000원", "date_start": "2026-04-01"},
        ],
    }

    output_path = tmp_path / "test-output.docx"
    generate_document(Path("io-sample.docx"), data, output_path)

    doc = Document(output_path)

    # Table 0 (Contact Info)에서 client_name 확인
    table0 = doc.tables[0]
    assert "테스트매체A" in table0.rows[1].cells[4].text

    # Table 2에서 gross_rate 확인
    table2 = doc.tables[2]
    assert "50%" in table2.rows[0].cells[1].text

    # Table 3 (서명)에서 client_manager 확인
    table3 = doc.tables[3]
    assert "김철수" in table3.rows[1].cells[3].text


def test_dynamic_widget_rows(tmp_path):
    """위젯 수에 따라 Table 1의 행이 동적으로 생성된다."""
    from generate import generate_document

    data = {
        "client_name": "테스트매체",
        "client_address": "주소",
        "client_email": "e@e.com",
        "client_manager": "매니저",
        "gross_rate": "40%",
        "widgets": [
            {"service": f"매체{i}", "service_name": f"s{i}.com",
             "widget_name": f"위젯{i}", "value": f"CPM {i}00원",
             "date_start": "2026-04-01"}
            for i in range(5)
        ],
    }

    output_path = tmp_path / "test-dynamic.docx"
    generate_document(Path("io-sample.docx"), data, output_path)

    doc = Document(output_path)
    table1 = doc.tables[1]

    # 헤더 1행 + 위젯 5행 = 6행
    assert len(table1.rows) == 6

    # 각 행의 첫 번째 셀에 매체명이 들어있는지 확인
    for i in range(5):
        assert f"매체{i}" in table1.rows[i + 1].cells[0].text
